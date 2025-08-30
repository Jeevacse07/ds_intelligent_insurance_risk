import streamlit as st
import pandas as pd
import joblib
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline, AutoModelForCausalLM
from docx import Document
import PyPDF2
import fitz  # PyMuPDF
import re
from summa.summarizer import summarize
from pathlib import Path

# =====================================================
# Load Light Models (fast, no caching needed)
# =====================================================
rf_high_risk = joblib.load("models/rf_high_risk_predict_model.joblib")
rf_claim_amount = joblib.load("models/rf_predict_claim_model.pkl")

# =====================================================
# Cached Model Loaders
# =====================================================
@st.cache_resource
def load_translation_model():
    model_name = "facebook/m2m100_418M"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)
    return tokenizer, model, device

@st.cache_resource
def load_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

@st.cache_resource
def load_qa_model():
    model_path = Path("D:/Jeeva/ds_course/insurance_final_project/app/insurance-gpt")
    model = AutoModelForCausalLM.from_pretrained(str(model_path))
    tokenizer = AutoTokenizer.from_pretrained(str(model_path))
    return model, tokenizer

# Load cached models once
tokenizer_translation, model_translation, device = load_translation_model()
abstractive_summarizer = load_summarizer()
qa_model, qa_tokenizer = load_qa_model()

# =====================================================
# Utility Functions
# =====================================================
def read_docx(file_path):
    doc = Document(file_path)
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]

def read_pdf(file_path):
    pdf_reader = PyPDF2.PdfReader(file_path)
    texts = []
    for page in pdf_reader.pages:
        texts.append(page.extract_text())
    return [t.strip() for t in texts if t.strip()]

def read_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return [line.strip() for line in f if line.strip()]

def translate_batch(texts, target_lang="ta", batch_size=8):
    tokenizer_translation.src_lang = "en"
    translated_texts = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i+batch_size]
        inputs = tokenizer_translation(batch, return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
        generated_tokens = model_translation.generate(
            **inputs,
            forced_bos_token_id=tokenizer_translation.get_lang_id(target_lang),
            max_length=600
        )
        translations = tokenizer_translation.batch_decode(generated_tokens, skip_special_tokens=True)
        translated_texts.extend(translations)
    return translated_texts

def translate_document(file_path, target_languages=["ta", "fr", "es"], file_type="docx"):
    if file_type == "docx":
        texts = read_docx(file_path)
    elif file_type == "pdf":
        texts = read_pdf(file_path)
    elif file_type == "txt":
        texts = read_txt(file_path)
    else:
        raise ValueError("Unsupported file type")

    translated_docs = {}
    for lang in target_languages:
        translated_docs[lang] = translate_batch(texts, target_lang=lang)
    return translated_docs

def chunk_text(text, max_sentences=50):
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = [' '.join(sentences[i:i+max_sentences]) for i in range(0, len(sentences), max_sentences)]
    return chunks

def extract_text(file_path):
    if file_path.lower().endswith('.pdf'):
        doc = fitz.open(file_path)
        text = "".join([page.get_text() for page in doc])
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
    return text

def summarize_insurance_doc(file_path, extract_ratio=0.1, max_abs_len=200, min_abs_len=50, chunk_size=50):
    text = extract_text(file_path)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^a-zA-Z0-9.,!? ]', '', text)

    text_chunks = chunk_text(text, max_sentences=chunk_size)

    extractive_summary_chunks = []
    for chunk in text_chunks:
        summary = summarize(chunk, ratio=extract_ratio)
        if not summary.strip():
            summary = chunk
        extractive_summary_chunks.append(summary)
    extractive_summary = ' '.join(extractive_summary_chunks)

    abs_chunks = chunk_text(extractive_summary, max_sentences=20)
    abstractive_summary_list = [
        abstractive_summarizer(chunk, max_length=max_abs_len, min_length=min_abs_len, do_sample=False)[0]['summary_text']
        for chunk in abs_chunks
    ]
    abstractive_summary = ' '.join(abstractive_summary_list)

    return extractive_summary, abstractive_summary

def generate_response(query, max_length=150):
    input_ids = qa_tokenizer(f"Question: {query} Answer:", return_tensors="pt").input_ids
    output = qa_model.generate(
        input_ids,
        max_length=max_length,
        num_beams=5,
        no_repeat_ngram_size=2,
        early_stopping=True,
        pad_token_id=qa_tokenizer.eos_token_id,
        temperature=0.7,
        do_sample=True
    )
    return qa_tokenizer.decode(output[0], skip_special_tokens=True)

# =====================================================
# Streamlit App
# =====================================================
st.sidebar.title("Insurance Prediction Modules")
module = st.sidebar.selectbox(
    "Choose Module",
    [
        "High-Risk Claim Prediction",
        "Claim Amount Prediction",
        "Multilingual Insurance Document Translation",
        "Insurance Document Summary",
        "Automated Insurance Response"
    ]
)

# ---------------- Module 1: High-Risk Prediction ----------------
if module == "High-Risk Claim Prediction":
    st.title("High-Risk Insurance Prediction Module")
    police_report = st.selectbox("Police Report Filed?", ["No", "Yes"])
    police_report = 1 if police_report == "Yes" else 0
    total_policy_claims = st.number_input("Total Policy Claims", min_value=1, max_value=10, value=1)
    incident_cause = st.selectbox("Incident Cause", ["Driver Error", "Natural Causes", "Other Causes", "Other Driver Error"])
    claim_type = st.selectbox("Claim Type", ["Material and Injury", "Material Only"])
    claim_area_home = st.checkbox("Claim Area: Home")
    claim_date = st.date_input("Claim Date")

    input_data = pd.DataFrame({
        'police_report': [police_report],
        'total_policy_claims': [total_policy_claims],
        'incident_cause_driver_error': [1 if incident_cause=="Driver Error" else 0],
        'incident_cause_natural_causes': [1 if incident_cause=="Natural Causes" else 0],
        'incident_cause_other_causes': [1 if incident_cause=="Other Causes" else 0],
        'incident_cause_other_driver_error': [1 if incident_cause=="Other Driver Error" else 0],
        'claim_area_home': [int(claim_area_home)],
        'claim_type_material_and_injury': [1 if claim_type=="Material and Injury" else 0],
        'claim_type_material_only': [1 if claim_type=="Material Only" else 0],
        'claim_year': [claim_date.year],
        'claim_month': [claim_date.month],
        'claim_day': [claim_date.day]
    })

    if st.button("Predict High-Risk Claim"):
        prediction = rf_high_risk.predict(input_data)
        result = "High-Risk" if prediction[0] == 1 else "Low-Risk"
        st.subheader(f"Prediction Result: {result}")

# ---------------- Module 2: Claim Amount Prediction ----------------
elif module == "Claim Amount Prediction":
    st.title("Claim Amount Prediction")
    police_report = st.selectbox("Police Report Filed?", ["No", "Yes"], key="claim_police")
    police_report = 1 if police_report == "Yes" else 0
    total_policy_claims = st.number_input("Total Policy Claims", min_value=1, max_value=10, value=1, key="claim_total_claims")
    incident_cause = st.selectbox("Incident Cause", ["Driver Error", "Natural Causes", "Other Causes", "Other Driver Error"], key="claim_incident")
    claim_type_material_and_injury = st.checkbox("Claim Type: Material and Injury", key="claim_type1")
    claim_area_home = st.checkbox("Claim Area: Home", key="claim_area")
    claim_date = st.date_input("Claim Date", key="claim_date")
    fraudulent = st.selectbox("Fraudulent Claim?", ["No", "Yes"], key="fraudulent")
    fraudulent = 1 if fraudulent == "Yes" else 0
    high_risk = st.selectbox("High-Risk Claim?", ["No", "Yes"], key="high_risk")
    high_risk = 1 if high_risk == "Yes" else 0

    input_data_claim = pd.DataFrame({
        'police_report': [police_report],
        'total_policy_claims': [total_policy_claims],
        'fraudulent': [fraudulent],
        'incident_cause_driver_error': [1 if incident_cause=="Driver Error" else 0],
        'incident_cause_natural_causes': [1 if incident_cause=="Natural Causes" else 0],
        'incident_cause_other_causes': [1 if incident_cause=="Other Causes" else 0],
        'incident_cause_other_driver_error': [1 if incident_cause=="Other Driver Error" else 0],
        'claim_area_home': [int(claim_area_home)],
        'claim_type_material_and_injury': [int(claim_type_material_and_injury)],
        'claim_year': [claim_date.year],
        'claim_month': [claim_date.month],
        'claim_day': [claim_date.day],
        'high_risk': [high_risk]
    })

    trained_features = rf_claim_amount.feature_names_in_
    for col in trained_features:
        if col not in input_data_claim.columns:
            input_data_claim[col] = 0
    input_data_claim = input_data_claim[trained_features]

    if st.button("Predict Claim Amount"):
        prediction = rf_claim_amount.predict(input_data_claim)
        st.subheader(f"Predicted Claim Amount: ₹{prediction[0]:,.2f}")

# ---------------- Module 3: Multilingual Translation ----------------
elif module == "Multilingual Insurance Document Translation":
    st.title("Multilingual Insurance Document Translation")
    uploaded_file = st.file_uploader("Upload Insurance Document", type=["docx","pdf","txt"])
    target_langs = st.multiselect("Select Target Languages", ["ta","fr","es","hi","de"], default=["ta","fr","es"])
    if uploaded_file is not None:
        file_type = uploaded_file.name.split(".")[-1].lower()
        if st.button("Translate Document"):
            with open("temp."+file_type, "wb") as f:
                f.write(uploaded_file.getbuffer())
            translations = translate_document("temp."+file_type, target_languages=target_langs, file_type=file_type)
            for lang, text in translations.items():
                st.subheader(f"Translation ({lang})")
                st.text_area("", "\n".join(text), height=200)

# ---------------- Module 4: Insurance Summary ----------------
elif module == "Insurance Document Summary":
    st.title("Insurance Document Summarization")
    uploaded_file = st.file_uploader("Upload Insurance Document", type=["docx","pdf"])
    if uploaded_file is not None:
        file_type = uploaded_file.name.split(".")[-1].lower()
        if st.button("Summarize Document"):
            with open("temp."+file_type, "wb") as f:
                f.write(uploaded_file.getbuffer())
            extractive, abstractive = summarize_insurance_doc("temp."+file_type)
            st.subheader("Extractive Summary")
            st.write(extractive)
            st.subheader("Abstractive Summary")
            st.write(abstractive)

# ---------------- Module 5: Automated Insurance Q&A ----------------
elif module == "Automated Insurance Response":
    st.title("Automated Insurance Q&A")
    user_query = st.text_input("Ask an insurance-related question")
    if st.button("Get Response") and user_query:
        response = generate_response(user_query)
        if "Answer:" in response:
            answer = response.split("Answer:")[-1].strip()
        else:
            answer = response
        st.subheader("AI Response")
        st.write(answer)
